#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
源代码导出脚本
将src文件夹下的所有源代码文件导出到Word文档中
每个文件的源码前会加上文件名标识
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches
from datetime import datetime

def should_include_file(file_path):
    """
    判断是否应该包含该文件
    排除图片、音频、视频、二进制文件等不需要的文件类型
    """
    # 需要排除的文件扩展名
    exclude_extensions = {
        '.png', '.jpg', '.jpeg', '.gif', '.bmp', '.ico', '.svg',
        '.mp3', '.mp4', '.avi', '.mov', '.wmv', '.flv',
        '.zip', '.rar', '.7z', '.tar', '.gz',
        '.pdf', '.doc', '.docx', '.xls', '.xlsx',
        '.ttf', '.woff', '.woff2',
        '.pyc', '.pyo', '.pyd',
        '.log', '.tmp', '.cache'
    }

    # 需要包含的文件扩展名
    include_extensions = {
        '.js', '.vue', '.ts', '.tsx', '.jsx',
        '.html', '.htm', '.css', '.scss', '.less',
        '.json', '.xml', '.md', '.txt',
        '.py', '.java', '.c', '.cpp', '.h',
        '.php', '.rb', '.go', '.rs',
        '.sql', '.sh', '.bat', '.ps1'
    }

    file_extension = file_path.suffix.lower()

    # 如果在排除列表中，直接排除
    if file_extension in exclude_extensions:
        return False

    # 如果在包含列表中，包含
    if file_extension in include_extensions:
        return True

    # 对于没有扩展名的文件，检查是否是可读文本文件
    if not file_extension:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                f.read(1024)  # 尝试读取前1024个字符
            return True
        except (UnicodeDecodeError, PermissionError):
            return False

    # 对于其他扩展名，默认排除
    return False

def get_relative_path(file_path, base_path):
    """
    获取文件的相对路径
    """
    try:
        relative_path = file_path.relative_to(base_path)
        return str(relative_path)
    except ValueError:
        return str(file_path)

def read_file_content(file_path):
    """
    读取文件内容
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return content
    except (UnicodeDecodeError, PermissionError) as e:
        print(f"警告: 无法读取文件 {file_path}: {e}")
        return f"无法读取文件内容: {e}"
    except Exception as e:
        print(f"错误: 读取文件 {file_path} 时发生未知错误: {e}")
        return f"读取文件时发生错误: {e}"

def export_source_code(src_path, output_path):
    """
    导出源代码到Word文档
    """
    src_path = Path(src_path)
    output_path = Path(output_path)

    if not src_path.exists():
        print(f"错误: 源文件夹不存在 - {src_path}")
        return False

    # 创建输出目录（如果不存在）
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # 创建Word文档
    doc = Document()
    doc.add_heading('源代码文档', 0)
    doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'源文件夹: {src_path.absolute()}')
    doc.add_paragraph('')

    # 收集所有需要处理的文件
    files_to_process = []

    print("正在扫描文件...")
    for root, dirs, files in os.walk(src_path):
        # 跳过node_modules, .git, dist等目录
        dirs[:] = [d for d in dirs if d not in {'node_modules', '.git', 'dist', 'build', '__pycache__'}]

        for file in files:
            file_path = Path(root) / file
            if should_include_file(file_path):
                files_to_process.append(file_path)

    print(f"找到 {len(files_to_process)} 个文件需要处理")

    # 处理每个文件
    processed_count = 0
    for file_path in sorted(files_to_process):
        try:
            relative_path = get_relative_path(file_path, src_path)
            content = read_file_content(file_path)

            # 添加文件标题
            doc.add_heading(f'文件: {relative_path}', level=1)

            # 添加文件信息
            file_info = f"文件路径: {file_path.absolute()}\n"
            file_info += f"文件大小: {file_path.stat().st_size} 字节\n"
            file_info += f"修改时间: {datetime.fromtimestamp(file_path.stat().st_mtime).strftime('%Y-%m-%d %H:%M:%S')}\n"
            doc.add_paragraph(file_info)

            # 添加代码内容
            if content.strip():
                # 创建代码段落
                code_paragraph = doc.add_paragraph()
                code_paragraph.add_run(content)
                # 设置等宽字体以保持代码格式
                for run in code_paragraph.runs:
                    run.font.name = 'Courier New'
                    run.font.size = 100000  # 10pt in twentieths of a point
            else:
                doc.add_paragraph("(文件为空或无法读取)")

            # 添加分隔线
            doc.add_paragraph('=' * 50)
            doc.add_paragraph('')

            processed_count += 1
            if processed_count % 10 == 0:
                print(f"已处理 {processed_count} 个文件...")

        except Exception as e:
            print(f"处理文件 {file_path} 时发生错误: {e}")
            doc.add_heading(f'文件: {relative_path} (处理出错)', level=1)
            doc.add_paragraph(f"错误信息: {e}")
            doc.add_paragraph('=' * 50)
            doc.add_paragraph('')

    # 添加统计信息
    doc.add_heading('统计信息', level=1)
    doc.add_paragraph(f'总文件数: {len(files_to_process)}')
    doc.add_paragraph(f'成功处理: {processed_count}')
    doc.add_paragraph(f'处理时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')

    # 保存文档
    try:
        doc.save(output_path)
        print(f"成功导出到: {output_path}")
        print(f"总共处理了 {processed_count} 个文件")
        return True
    except Exception as e:
        print(f"保存文档时发生错误: {e}")
        return False

def main():
    # 设置路径
    src_folder = "src"
    output_file = r"C:\Users\78623\OneDrive\文档\论文\郑老师基金\北邮软著登记准备材料-2025\4.源代码.doc"

    print("开始导出源代码...")
    print(f"源文件夹: {src_folder}")
    print(f"输出文件: {output_file}")
    print("-" * 50)

    success = export_source_code(src_folder, output_file)

    if success:
        print("\n导出完成!")
    else:
        print("\n导出失败!")
        sys.exit(1)

if __name__ == "__main__":
    main()
